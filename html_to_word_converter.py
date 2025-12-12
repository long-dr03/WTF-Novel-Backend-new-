#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
HTML to Word Converter Tool
Chuyển đổi các file HTML từ truyện web sang định dạng Word (.docx)

Author: Novel Converter
Version: 1.0.0
"""

import os
import re
import glob
import argparse
from pathlib import Path
from typing import List, Optional
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class HTMLToWordConverter:
    """
    Class chuyển đổi file HTML sang Word document
    """
    
    def __init__(
        self,
        font_name: str = "Times New Roman",
        font_size: int = 13,
        line_spacing: float = 1.5,
        chapter_font_size: int = 16,
        page_margin: float = 2.54  # cm
    ):
        """
        Khởi tạo converter với các tùy chọn định dạng
        
        Args:
            font_name: Tên font chữ (mặc định: Times New Roman)
            font_size: Cỡ chữ nội dung (mặc định: 13pt)
            line_spacing: Khoảng cách dòng (mặc định: 1.5)
            chapter_font_size: Cỡ chữ tiêu đề chương (mặc định: 16pt)
            page_margin: Lề trang tính bằng cm (mặc định: 2.54cm)
        """
        self.font_name = font_name
        self.font_size = font_size
        self.line_spacing = line_spacing
        self.chapter_font_size = chapter_font_size
        self.page_margin = page_margin
    
    def parse_html_file(self, html_path: str) -> dict:
        """
        Đọc và phân tích file HTML
        
        Args:
            html_path: Đường dẫn đến file HTML
            
        Returns:
            Dictionary chứa title và content
        """
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Lấy tiêu đề chương từ tag h1 hoặc title
        title = ""
        h1_tag = soup.find('h1')
        if h1_tag:
            title = h1_tag.get_text(strip=True)
        else:
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text(strip=True)
        
        # Lấy nội dung từ div#content-chapter
        content_div = soup.find('div', id='content-chapter')
        
        paragraphs = []
        if content_div:
            # Lấy tất cả các thẻ p trong content div
            p_tags = content_div.find_all('p')
            for p in p_tags:
                text = p.get_text(strip=True)
                if text:  # Chỉ thêm đoạn văn không rỗng
                    paragraphs.append(text)
        else:
            # Fallback: lấy tất cả các thẻ p trong body
            body = soup.find('body')
            if body:
                p_tags = body.find_all('p')
                for p in p_tags:
                    text = p.get_text(strip=True)
                    if text:
                        paragraphs.append(text)
        
        return {
            'title': title,
            'paragraphs': paragraphs
        }
    
    def set_cell_font(self, run, font_name: str):
        """
        Set font cho cả tiếng Việt và các ký tự khác
        """
        run.font.name = font_name
        # Đặt font cho East Asian characters
        r = run._element
        rFonts = r.find(qn('w:rPr'))
        if rFonts is None:
            rFonts = OxmlElement('w:rPr')
            r.insert(0, rFonts)
        fonts = rFonts.find(qn('w:rFonts'))
        if fonts is None:
            fonts = OxmlElement('w:rFonts')
            rFonts.insert(0, fonts)
        fonts.set(qn('w:eastAsia'), font_name)
        fonts.set(qn('w:ascii'), font_name)
        fonts.set(qn('w:hAnsi'), font_name)
    
    def create_document(self) -> Document:
        """
        Tạo document Word mới với các style được cấu hình
        
        Returns:
            Document object
        """
        doc = Document()
        
        # Thiết lập lề trang
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(self.page_margin)
            section.bottom_margin = Cm(self.page_margin)
            section.left_margin = Cm(self.page_margin)
            section.right_margin = Cm(self.page_margin)
        
        return doc
    
    def add_chapter_to_document(
        self, 
        doc: Document, 
        chapter_data: dict,
        add_page_break: bool = True
    ):
        """
        Thêm một chương vào document
        
        Args:
            doc: Document object
            chapter_data: Dictionary chứa title và paragraphs
            add_page_break: Có thêm page break trước chương không
        """
        # Thêm page break nếu cần (không phải chương đầu)
        if add_page_break and len(doc.paragraphs) > 0:
            doc.add_page_break()
        
        # Thêm tiêu đề chương
        if chapter_data['title']:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(chapter_data['title'])
            title_run.bold = True
            title_run.font.size = Pt(self.chapter_font_size)
            self.set_cell_font(title_run, self.font_name)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.space_after = Pt(12)
        
        # Thêm nội dung
        for text in chapter_data['paragraphs']:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(self.font_size)
            self.set_cell_font(run, self.font_name)
            
            # Thiết lập paragraph format
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Cm(1.27)  # Thụt đầu dòng
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = self.line_spacing
    
    def convert_single_file(
        self, 
        html_path: str, 
        output_path: Optional[str] = None
    ) -> str:
        """
        Chuyển đổi một file HTML sang Word
        
        Args:
            html_path: Đường dẫn file HTML nguồn
            output_path: Đường dẫn file Word đích (tùy chọn)
            
        Returns:
            Đường dẫn file Word đã tạo
        """
        if not output_path:
            output_path = str(Path(html_path).with_suffix('.docx'))
        
        chapter_data = self.parse_html_file(html_path)
        doc = self.create_document()
        self.add_chapter_to_document(doc, chapter_data, add_page_break=False)
        doc.save(output_path)
        
        return output_path
    
    def convert_multiple_files(
        self,
        html_paths: List[str],
        output_path: str,
        sort_files: bool = True
    ) -> str:
        """
        Chuyển đổi nhiều file HTML thành một file Word
        
        Args:
            html_paths: Danh sách đường dẫn file HTML
            output_path: Đường dẫn file Word đích
            sort_files: Có sắp xếp file theo tên không
            
        Returns:
            Đường dẫn file Word đã tạo
        """
        if sort_files:
            html_paths = sorted(html_paths, key=lambda x: self._extract_chapter_number(x))
        
        doc = self.create_document()
        
        for i, html_path in enumerate(html_paths):
            print(f"Đang xử lý: {os.path.basename(html_path)}")
            chapter_data = self.parse_html_file(html_path)
            self.add_chapter_to_document(
                doc, 
                chapter_data, 
                add_page_break=(i > 0)  # Page break từ chương 2 trở đi
            )
        
        doc.save(output_path)
        return output_path
    
    def convert_folder(
        self,
        folder_path: str,
        output_path: Optional[str] = None,
        pattern: str = "*.html",
        combine: bool = True
    ) -> str:
        """
        Chuyển đổi tất cả file HTML trong một thư mục
        
        Args:
            folder_path: Đường dẫn thư mục chứa file HTML
            output_path: Đường dẫn file/thư mục đích
            pattern: Pattern để match file (mặc định: *.html)
            combine: Gộp tất cả thành một file Word hay không
            
        Returns:
            Đường dẫn file/thư mục Word đã tạo
        """
        # Tìm tất cả file HTML
        html_files = glob.glob(os.path.join(folder_path, pattern))
        
        if not html_files:
            raise ValueError(f"Không tìm thấy file HTML nào trong {folder_path}")
        
        print(f"Tìm thấy {len(html_files)} file HTML")
        
        if combine:
            # Gộp tất cả vào một file
            if not output_path:
                folder_name = os.path.basename(folder_path.rstrip('/\\'))
                output_path = os.path.join(folder_path, f"{folder_name}_combined.docx")
            
            return self.convert_multiple_files(html_files, output_path)
        else:
            # Chuyển từng file riêng biệt
            if not output_path:
                output_path = folder_path
            
            os.makedirs(output_path, exist_ok=True)
            
            for html_path in sorted(html_files, key=lambda x: self._extract_chapter_number(x)):
                filename = Path(html_path).stem + '.docx'
                docx_path = os.path.join(output_path, filename)
                self.convert_single_file(html_path, docx_path)
                print(f"Đã tạo: {filename}")
            
            return output_path
    
    @staticmethod
    def _extract_chapter_number(filename: str) -> int:
        """
        Trích xuất số chương từ tên file để sắp xếp
        
        Args:
            filename: Tên file
            
        Returns:
            Số chương (hoặc 0 nếu không tìm thấy)
        """
        basename = os.path.basename(filename)
        # Tìm số trong tên file (ví dụ: chuong-001.html -> 1)
        match = re.search(r'(\d+)', basename)
        if match:
            return int(match.group(1))
        return 0


def main():
    """
    Main function - xử lý command line arguments
    """
    parser = argparse.ArgumentParser(
        description='Chuyển đổi file HTML từ truyện web sang Word',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ví dụ sử dụng:
  # Chuyển tất cả file HTML trong thư mục thành 1 file Word
  python html_to_word_converter.py -i "example - content" -o "truyen.docx"
  
  # Chuyển từng file riêng biệt
  python html_to_word_converter.py -i "example - content" --separate
  
  # Chuyển một file HTML
  python html_to_word_converter.py -i "chuong-001.html" -o "chuong-001.docx"
  
  # Tùy chỉnh định dạng
  python html_to_word_converter.py -i "example - content" -o "truyen.docx" --font-size 14 --font "Arial"
        """
    )
    
    parser.add_argument(
        '-i', '--input',
        required=True,
        help='Đường dẫn file HTML hoặc thư mục chứa file HTML'
    )
    
    parser.add_argument(
        '-o', '--output',
        help='Đường dẫn file Word đầu ra (mặc định: tự động tạo)'
    )
    
    parser.add_argument(
        '--separate',
        action='store_true',
        help='Chuyển từng file HTML thành file Word riêng biệt'
    )
    
    parser.add_argument(
        '--font',
        default='Times New Roman',
        help='Tên font chữ (mặc định: Times New Roman)'
    )
    
    parser.add_argument(
        '--font-size',
        type=int,
        default=13,
        help='Cỡ chữ nội dung (mặc định: 13)'
    )
    
    parser.add_argument(
        '--chapter-font-size',
        type=int,
        default=16,
        help='Cỡ chữ tiêu đề chương (mặc định: 16)'
    )
    
    parser.add_argument(
        '--line-spacing',
        type=float,
        default=1.5,
        help='Khoảng cách dòng (mặc định: 1.5)'
    )
    
    parser.add_argument(
        '--margin',
        type=float,
        default=2.54,
        help='Lề trang tính bằng cm (mặc định: 2.54)'
    )
    
    parser.add_argument(
        '--pattern',
        default='*.html',
        help='Pattern để match file HTML (mặc định: *.html)'
    )
    
    args = parser.parse_args()
    
    # Tạo converter với các tùy chọn
    converter = HTMLToWordConverter(
        font_name=args.font,
        font_size=args.font_size,
        chapter_font_size=args.chapter_font_size,
        line_spacing=args.line_spacing,
        page_margin=args.margin
    )
    
    input_path = args.input
    output_path = args.output
    
    try:
        if os.path.isfile(input_path):
            # Chuyển một file
            result = converter.convert_single_file(input_path, output_path)
            print(f"✓ Đã tạo file: {result}")
        elif os.path.isdir(input_path):
            # Chuyển thư mục
            result = converter.convert_folder(
                input_path,
                output_path,
                pattern=args.pattern,
                combine=not args.separate
            )
            if args.separate:
                print(f"✓ Đã tạo các file Word trong: {result}")
            else:
                print(f"✓ Đã tạo file: {result}")
        else:
            print(f"✗ Lỗi: Không tìm thấy '{input_path}'")
            return 1
            
    except Exception as e:
        print(f"✗ Lỗi: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())
